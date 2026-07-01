from copy import deepcopy
from pathlib import Path
import shutil
import zipfile

from docx import Document
from docx.enum.text import WD_COLOR_INDEX


SRC = Path(r"C:\Users\jenny.lu\Documents\建大輪胎專案\期中交付文件\DOC-02功能規格書畫面規格書2.0.docx")
OUT = Path(r"C:\Users\jenny.lu\Documents\建大輪胎專案\期中交付文件\DOC-02功能規格書畫面規格書3.0.docx")


def clone_run_style(src, dst):
    dst.bold = src.bold
    dst.italic = src.italic
    dst.underline = src.underline
    if src.font is not None:
        dst.font.name = src.font.name
        dst.font.size = src.font.size
        dst.font.color.rgb = src.font.color.rgb


def set_paragraph_text(paragraph, text):
    base = paragraph.runs[0] if paragraph.runs else None
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    if base is not None:
        clone_run_style(base, run)


def set_cell_text(cell, text):
    paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, text)
    for extra in cell.paragraphs[1:]:
        set_paragraph_text(extra, "")


def append_highlight(paragraph, text):
    if paragraph.text and not paragraph.text.endswith((" ", "\n")):
        paragraph.add_run("\n")
    run = paragraph.add_run(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True
    return run


def append_highlight_to_cell(cell, text):
    append_highlight(cell.paragraphs[-1], text)


def replace_in_paragraph(paragraph, old, new):
    if old not in paragraph.text:
        return False
    set_paragraph_text(paragraph, paragraph.text.replace(old, new))
    return True


def add_table_row_after(table, after_idx, values):
    tr = deepcopy(table.rows[after_idx]._tr)
    table._tbl.insert(after_idx + 1, tr)
    row = table.rows[after_idx + 1]
    for idx, value in enumerate(values):
        if idx < len(row.cells):
            set_cell_text(row.cells[idx], value)
    return row


def delete_table(table):
    tbl = table._element
    tbl.getparent().remove(tbl)


def patch_header_footer_versions(doc):
    for section in doc.sections:
        for part in (section.header, section.footer, section.first_page_header, section.first_page_footer):
            for paragraph in part.paragraphs:
                replace_in_paragraph(paragraph, "v1.2", "v3.0")
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_in_paragraph(paragraph, "v1.2", "v3.0")


def rebuild_appendix(doc):
    app = doc.tables[84]
    tbl = app._tbl
    for tr in list(tbl.tr_lst)[1:]:
        tbl.remove(tr)

    rows = [
        ["B-01", "文件資訊／品牌", "品牌名稱應為「建大輪胎（KENDA Tires）」；終端品牌=建大輪胎（KENDA Tires）。", "已統一品牌名稱為「建大輪胎（KENDA Tires）」。", "✓ 已修訂"],
        ["B-02", "§4.2 商品總覽", "下單功能：報價單預設「一般報價」，未來若增加「伙伴雲報價單」不顯示。", "【待確認／會議討論】伙伴雲報價單目前不預設顯示；需確認何時啟用、是否納入本版範圍，以及後台應如何設定顯示／不顯示規則。", "■ 待確認"],
        ["B-03", "§4.2 商品總覽", "登入即由帳號判斷客代所屬胎別，不需人工挑選；欄位應改為「產品系列」並定義分類與排序。", "欄位由「產品類別」改為「產品系列」，登入即依客代自動載入；補列 PCR／MC／BC 三大系列及子系列，與兩種客戶自定義排序規則。", "✓ 已修訂"],
        ["B-04", "§4.3 購物車", "希望交期空白時，轉檔以當月最後一日做為客戶希望交期。", "已於希望交期欄位載明預設規則，並刪除「可空白」字樣，避免誤解為無預設處理。", "✓ 已修訂"],
        ["B-05", "§4.3 購物車", "指配地點預設值為 01；需要轉出指配地點 01 代表的地址？", "已於指配地點欄位補充：下拉顯示「地點代號－地址」，預設值為 01；轉 ERP 時帶出該代號對應之實際送貨地址。", "✓ 已修訂"],
        ["B-06", "§4.3／拆單", "無三層報價功能需求；購物車轉 ERP 訂單拆單邏輯=報價單編號+包裝部門。", "已修正拆單口徑：本案無三層報價功能需求；購物車轉 ERP 訂單時，拆單邏輯以「報價單編號＋包裝部門」為準。", "✓ 已修訂"],
        ["B-07", "§4.6 獎金查詢", "獎金：1.固定：月/季/年獎金 2.販促：季度通案+個案 3.其它獎金。", "【待確認／會議討論】需確認獎金分類是否完整，以及是否需要對應後台設定畫面；本版先以螢光標示待確認。", "■ 待確認"],
        ["B-08", "§4.6 獎金查詢", "建議改為起迄年月，顯示符合期間的各項目總合。", "已將獎金查詢條件調整為起迄年月，並說明畫面應顯示符合期間之各項目總合。", "✓ 已修訂"],
        ["B-09", "店招畫面", "格式／項目與商企課謝雯萱規劃畫面不一致。", "【待確認／會議討論】需請建大內部先對齊商企課規劃版本與本專案規格基準，再回饋統一畫面規格。", "■ 待確認"],
        ["B-10", "帳號管理／付款方式", "要顯示 PD000 代表的意義；除顯示代碼外，也要同步顯示代碼說明。", "已於付款方式欄位補充：畫面須同步顯示代碼與代碼說明，例如「PD000－代碼說明」。", "✓ 已修訂"],
    ]
    for values in rows:
        row = app.add_row()
        for idx, value in enumerate(values):
            set_cell_text(row.cells[idx], value)
            if value.startswith("【待確認"):
                row.cells[idx].paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
                row.cells[idx].paragraphs[0].runs[0].bold = True


doc = Document(SRC)

# Document version and metadata.
replace_in_paragraph(doc.paragraphs[3], "KENDA B2B v1.2", "KENDA B2B v3.0")
replace_in_paragraph(doc.paragraphs[24], "文件版本\tv1.2", "文件版本\tv3.0")
set_paragraph_text(doc.paragraphs[25], "發行日期\t2026-06-22")
replace_in_paragraph(doc.paragraphs[30], "v1.1 → v1.2，2026-06-09", "v1.2 → v3.0，2026-06-22")
for paragraph in doc.paragraphs:
    replace_in_paragraph(paragraph, "v1.1 → v1.2，2026-06-09", "v1.2 → v3.0，2026-06-22")

info = doc.tables[0]
set_cell_text(info.rows[6].cells[1], "建大輪胎（KENDA Tires）")
set_cell_text(info.rows[9].cells[1], "v3.0")
set_cell_text(info.rows[10].cells[1], "2026-06-22")
set_cell_text(info.rows[11].cells[1], "依 0618 期中待修正文件與確認回覆修訂")

rev = doc.tables[1]
rev_row = rev.add_row()
for idx, value in enumerate(["v3.0", "2026-\n06-22", "艾創點", "依 0618 期中修訂清單更新：伙伴雲報價單待確認、指配地點地址、獎金查詢起迄年月、付款方式代碼說明、海外工廠範圍與商企課畫面差異待確認。"]):
    set_cell_text(rev_row.cells[idx], value)

# Shop / quote purpose.
shop = doc.tables[9]
set_cell_text(shop.rows[4].cells[1], "已登入（Bearer Token 有效）。報價用途預設為「一般報價」，需要其它報價單時再手動挑選；切換用途後 API 依 purpose 取對應報價價格。")
append_highlight_to_cell(shop.rows[4].cells[1], "【待確認／會議討論】伙伴雲報價單目前不預設顯示；需確認何時啟用、是否納入本版範圍，以及後台應如何設定顯示／不顯示規則。")

purpose = doc.tables[11]
set_cell_text(purpose.rows[1].cells[3], "一般報價／呆貨／販促；預設為一般報價，選擇後 API 依 purpose 取對應報價價格。")
append_highlight_to_cell(purpose.rows[1].cells[3], "【待確認／會議討論】伙伴雲報價單若未來啟用，需確認後台設定方式與前台顯示條件。")

products_api = doc.tables[13]
set_cell_text(products_api.rows[3].cells[3], "{ success, data: [ { quote_no, purpose, customer_code, lines: [ { quote_line_no, product_code, factory, unit_price, waste_tire_fee, pack_code, pack_qty } ] } ], meta: { total } }；伙伴雲報價單目前不預設顯示，是否啟用及後台設定方式待確認。")

standard = doc.tables[73]
set_cell_text(standard.rows[1].cells[1], "一般報價")

# Cart / delivery location / requested date.
cart_fields = doc.tables[17]
set_cell_text(cart_fields.rows[2].cells[3], "由 GET /shipping-locations 取得可選地點；格式：地點代號 — 地址。預設值為「01」。轉 ERP 時須帶出該指配地點代號對應之實際送貨地址。")
set_cell_text(cart_fields.rows[3].cells[4], "格式 YYYY-MM-DD；若未維護或未選擇，轉檔時預設以「當月最後一日」做為客戶希望交期。")

shipping_api = doc.tables[19]
set_cell_text(shipping_api.rows[1].cells[2], "{ success, data: [ { location_code, address, company_code } ] }；location_code=01 時亦須回傳對應之實際送貨地址。")

# Order query date range and split wording.
orders_fields = doc.tables[22]
set_cell_text(orders_fields.rows[1].cells[0], "start_ym／end_ym（起迄年月）")
set_cell_text(orders_fields.rows[1].cells[3], "格式 YYYY-MM（送 API 時轉換為 YYYYMM）；查詢符合起迄期間的訂單資料。")

split_result = doc.tables[83]
set_cell_text(split_result.rows[2].cells[1], "「本次訂單依『報價單編號＋包裝部門』規則，自動拆成 N 張建大訂單。」本案無三層報價功能需求，並已取消廢胎費拆單層。")
set_cell_text(split_result.rows[3].cells[1], "建大訂單號、報價單編號、包裝部門、包裝代號（含該語系包裝說明）、金額、品項數")

# Bonus query.
bonus = doc.tables[38]
set_cell_text(bonus.rows[3].cells[1], "依起迄年月查詢符合期間的獎金金額與各項目總合。獎金資料由建大 ERP 排程同步至 kenda.bonus。")
append_highlight_to_cell(bonus.rows[3].cells[1], "【待確認／會議討論】獎金分類是否採固定（月／季／年獎金）、販促（季度通案＋個案）、其它獎金，及是否需要後台設定畫面，需再確認。")

bonus_fields = doc.tables[39]
set_cell_text(bonus_fields.rows[1].cells[0], "start_ym／end_ym（起迄年月）")
set_cell_text(bonus_fields.rows[1].cells[3], "查詢起迄年月期間內符合條件之各項目總合。")
set_cell_text(bonus_fields.rows[2].cells[0], "period_label（顯示期間）")
set_cell_text(bonus_fields.rows[2].cells[3], "格式化顯示（例：2026-05～2026-08）")
add_table_row_after(bonus_fields, 3, ["bonus_category（獎金分類）", "唯讀／待確認", "—", "固定（月／季／年）、販促（季度通案＋個案）、其它獎金。", "【待確認／會議討論】分類完整性及後台設定方式待確認。"])
append_highlight_to_cell(bonus_fields.rows[4].cells[4], "")

bonus_api = doc.tables[42]
set_cell_text(bonus_api.rows[1].cells[2], "start_ym=YYYYMM / end_ym=YYYYMM / dealer=客代")
set_cell_text(bonus_api.rows[1].cells[3], "{ success, data: [ { customer_code, start_ym, end_ym, bonus_category?, amount } ], total? }")

# Signage screen conflict with customer's internal planning.
signage = doc.tables[52]
append_highlight_to_cell(signage.rows[3].cells[1], "【待確認／會議討論】此畫面格式與項目需與建大商企課謝雯萱規劃版本對齊後，再確認是否調整本規格。")

# Account payment code.
payment = doc.tables[69]
set_cell_text(payment.rows[4].cells[3], "例：PD000－代碼說明；畫面除顯示付款方式代碼外，須同步顯示該代碼代表之意義／說明。")

# Company / factory scope.
company = doc.tables[82]
set_cell_text(company.rows[1].cells[2], "KD（建大）；不含 KC/KT 等海外工廠")
append_highlight_to_cell(company.rows[1].cells[2], "【待確認／會議討論】海外工廠排除清單是否完整，需由建大確認。")

patch_header_footer_versions(doc)
rebuild_appendix(doc)

doc.save(OUT)

# Patch version text inside footer drawing/textbox XML when present.
tmp = OUT.with_suffix(".tmp.docx")
with zipfile.ZipFile(OUT, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename.startswith("word/footer") and item.filename.endswith(".xml"):
            data = data.replace(b"v1.2", b"v3.0")
        zout.writestr(item, data)
shutil.move(str(tmp), str(OUT))
print(str(OUT))
