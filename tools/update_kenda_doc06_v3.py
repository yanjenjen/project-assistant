from pathlib import Path
import shutil
import zipfile

from docx import Document
from docx.enum.text import WD_COLOR_INDEX


def find_source() -> Path:
    docs = Path.home() / "Documents"
    candidates = []
    for path in docs.rglob("DOC-06*.docx"):
        name = path.name
        if "2.0" in name and "(1)" not in name and "期中交付文件" in str(path):
            candidates.append(path)
    if not candidates:
        raise FileNotFoundError("DOC-06 2.0 source docx not found")
    return sorted(candidates, key=lambda p: p.stat().st_mtime, reverse=True)[0]


def clone_run_style(src, dst):
    dst.bold = src.bold
    dst.italic = src.italic
    dst.underline = src.underline
    if src.font is not None:
        dst.font.name = src.font.name
        dst.font.size = src.font.size
        dst.font.color.rgb = src.font.color.rgb


def set_paragraph_text(paragraph, text):
    base = paragraph.runs[0] if paragraph.runs else None
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    if base is not None:
        clone_run_style(base, run)


def replace_in_paragraph(paragraph, old, new):
    if old not in paragraph.text:
        return False
    set_paragraph_text(paragraph, paragraph.text.replace(old, new))
    return True


def append_highlight(paragraph, text):
    if paragraph.text and not paragraph.text.endswith(("\n", " ")):
        paragraph.add_run("\n")
    run = paragraph.add_run(text)
    run.bold = True
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return run


def patch_header_footer_versions(doc):
    for section in doc.sections:
        parts = (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
        )
        for part in parts:
            for paragraph in part.paragraphs:
                replace_in_paragraph(paragraph, "v1.2", "v3.0")
                replace_in_paragraph(paragraph, "2026-06-09", "2026-06-22")


src = find_source()
out = src.with_name(src.name.replace("2.0", "3.0"))

doc = Document(src)

# Cover metadata.
for paragraph in doc.paragraphs:
    replace_in_paragraph(paragraph, "版本：v1.2", "版本：v3.0")
    replace_in_paragraph(paragraph, "日期：2026-06-09", "日期：2026-06-22")
set_paragraph_text(doc.paragraphs[21], "專案：建大工業 KENDA B2B 經銷商平台版本：v3.0")
set_paragraph_text(doc.paragraphs[22], "日期：2026-06-22")

# Intro: DOC-06 page 1 asks whether non-ordering functions need fuller instructions.
append_highlight(
    doc.paragraphs[26],
    "【0618同步待確認】目前下單流程已有較完整步驟；訂單查詢、業績／獎金／對帳、最新消息、店招、合約下載、帳號管理等其他功能仍屬概要說明。是否需擴充為逐步操作手冊，需確認是否納入本次文件範圍。",
)

# Cart flow: clarify designated shipping location 01 and address display source.
set_paragraph_text(
    doc.paragraphs[33],
    "於商品卡輸入條數後「加入購物車」。進入購物車後選擇指配地點（預設值為 01，需同步顯示 01 代表的地址）與希望交期（建議 3 工作天以上）。",
)
set_paragraph_text(
    doc.paragraphs[34],
    "若未維護或未選擇希望交期，系統轉檔時預設以「當月最後一日」做為希望交期。可跨多張報價單、不同包裝的品項一起結帳。",
)

# Submit order: clarify split rules and package-code wording from DOC-06 page 4.
set_paragraph_text(
    doc.paragraphs[37],
    "按「確認送出訂單」。系統依「報價用途／報價單號 → 包裝單位」二層規則自動拆分成多張建大訂單。",
)
append_highlight(
    doc.paragraphs[37],
    "【0618同步待確認】拆單邏輯已排除廢胎費；是否原即不含廢胎費作為拆單依據，需明確書面確認。",
)
set_paragraph_text(
    doc.paragraphs[38],
    "每張訂單顯示建大訂單號、用途、報價單號、包裝代號及金額；包裝代號需同步顯示該語系的包裝說明，例如：裸裝＋1 張 Sticker。同一張 PO 拆出的子單歸於同一訂單群組（母子單）。",
)
append_highlight(
    doc.paragraphs[38],
    "【0618同步待確認】客戶註記之「36」指包裝代號，並非包裝部門；拆單依據與畫面顯示用語需統一確認，避免與 DOC-01/DOC-02 術語混淆。",
)

patch_header_footer_versions(doc)
doc.save(out)

# Patch version/date text inside footer/header drawing text boxes when present.
tmp = out.with_suffix(".tmp.docx")
with zipfile.ZipFile(out, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename.startswith(("word/header", "word/footer")) and item.filename.endswith(".xml"):
            data = data.replace(b"v1.2", b"v3.0")
            data = data.replace(b"2026-06-09", b"2026-06-22")
        zout.writestr(item, data)
shutil.move(str(tmp), str(out))
print(out)
