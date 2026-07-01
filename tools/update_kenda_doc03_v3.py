from pathlib import Path
import shutil
import zipfile

from docx import Document
from docx.enum.text import WD_COLOR_INDEX


SRC = Path(r"C:\Users\jenny.lu\Documents\建大輪胎專案\期中交付文件\DOC-03系統架構文件2.0.docx")
OUT = Path(r"C:\Users\jenny.lu\Documents\建大輪胎專案\期中交付文件\DOC-03系統架構文件3.0.docx")


def clone_run_style(src, dst):
    dst.bold = src.bold
    dst.italic = src.italic
    dst.underline = src.underline
    if src.font is not None:
        dst.font.name = src.font.name
        dst.font.size = src.font.size
        dst.font.color.rgb = src.font.color.rgb


def set_paragraph_text(paragraph, text):
    base = paragraph.runs[0] if paragraph.runs else None
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    if base is not None:
        clone_run_style(base, run)


def set_paragraph_plain_text(paragraph, text):
    for child in list(paragraph._p):
        if child.tag != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr":
            paragraph._p.remove(child)
    paragraph.add_run(text)


def set_cell_text(cell, text):
    paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, text)
    for extra in cell.paragraphs[1:]:
        set_paragraph_text(extra, "")


def append_highlight(paragraph, text):
    if paragraph.text and not paragraph.text.endswith((" ", "\n")):
        paragraph.add_run("\n")
    run = paragraph.add_run(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True
    return run


def append_highlight_to_cell(cell, text):
    append_highlight(cell.paragraphs[-1], text)


def replace_in_paragraph(paragraph, old, new):
    if old not in paragraph.text:
        return False
    set_paragraph_text(paragraph, paragraph.text.replace(old, new))
    return True


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def patch_header_footer_versions(doc):
    for section in doc.sections:
        for part in (section.header, section.footer, section.first_page_header, section.first_page_footer):
            for paragraph in part.paragraphs:
                replace_in_paragraph(paragraph, "v1.0", "v3.0")
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_in_paragraph(paragraph, "v1.0", "v3.0")


doc = Document(SRC)

# Version and metadata.
replace_in_paragraph(doc.paragraphs[3], "KENDA B2B v1.0", "KENDA B2B v3.0")
replace_in_paragraph(doc.paragraphs[24], "文件版本\tv1.0", "文件版本\tv3.0")
set_paragraph_text(doc.paragraphs[25], "發行日期\t2026-06-22")

info = doc.tables[0]
set_cell_text(info.rows[3].cells[1], "v3.0")
set_cell_text(info.rows[4].cells[1], "2026-06-22")
set_cell_text(info.rows[9].cells[1], "依 0618 期中待修正文件與確認口徑修訂")

rev = doc.tables[1]
row = rev.add_row()
for idx, value in enumerate(["v3.0", "2026-06-22", "艾創點", "依 0618 客戶註記修訂：公告分類邏輯、KD/KF 廢胎費規則、公司別＋胎別介面/語系判斷、拆單口徑與 V4 多公司別全球語系架構預留待確認。"]):
    set_cell_text(row.cells[idx], value)

# Frontend pages and modules.
pages = doc.tables[10]
set_cell_text(pages.rows[8].cells[2], "起迄年月獎金；分類呈現含固定（月／季／年）、販促（季度通案＋個案）、其它獎金。")
append_highlight_to_cell(pages.rows[8].cells[2], "【待確認／會議討論】需確認獎金分類完整性與是否需要後台設定畫面。")
set_cell_text(pages.rows[10].cells[2], "依公司別＋胎別業務人員推播之分類公告／最新消息，並非單純以公司別區分。")
set_cell_text(pages.rows[11].cells[0], "廣告資產申請")
set_cell_text(pages.rows[11].cells[2], "廣告資產申請及查詢。")
append_highlight_to_cell(pages.rows[11].cells[2], "【待確認／會議討論】需確認此為既有店招申請功能延伸，或屬新增廣告資產管理模組。")
set_cell_text(pages.rows[12].cells[0], "廣告資產合約下載")
set_cell_text(pages.rows[12].cells[2], "各式廣告資產合約 PDF 下載。")

models = doc.tables[11]
set_cell_text(models.rows[5].cells[2], "含一般／呆貨／販促三種用途；KD 單價不含廢胎費、KF 單價內含廢胎費。")
append_highlight_to_cell(models.rows[5].cells[2], "【待確認／會議討論】需確認 KD/KF 與公司別 1/F 對應口徑，以及廢胎費扣減規則是否完全一致。")
set_cell_text(models.rows[9].cells[2], "獎金暫存；分類呈現包含固定（月／季／年）、販促（季度通案＋個案）、其它獎金。")
set_cell_text(models.rows[12].cells[2], "依公司別＋胎別業務人員推播公告及最新消息；分類邏輯不以公司別作為唯一區分。")
set_cell_text(models.rows[13].cells[1], "廣告資產申請")
set_cell_text(models.rows[13].cells[2], "廣告資產申請／查詢流程；是否屬新增模組待確認。")
set_cell_text(models.rows[14].cells[1], "廣告資產合約下載")
set_cell_text(models.rows[14].cells[2], "各式廣告資產合約 PDF 儲存與下載。")

# Clean up broken inherited-model prose and remove the floating object that
# makes Word wrap the preceding line into narrow columns.
set_paragraph_plain_text(doc.paragraphs[94], "資料模型全新命名 kenda.*，不延伸原生 sale.order / product.template；僅 kenda.dealer 與 kenda.account 採委託繼承，分別對應 res.partner 與 res.users。")
for idx in (95, 96, 97, 98, 99, 100, 101, 102):
    set_paragraph_plain_text(doc.paragraphs[idx], "")
set_paragraph_plain_text(doc.paragraphs[103], "所有欄位以 Excel 12 張表（下單平台資料交換260522.xlsx）為唯一真實來源（SoT）。")

# ERP and performance logic.
set_paragraph_text(doc.paragraphs[121], "業績計算原則：")
set_paragraph_text(doc.paragraphs[122], "業績由 ERP 端依出貨金額及廢胎費規則計算後回傳，平台僅接收最終值與業績明細子陣列。")
set_paragraph_text(doc.paragraphs[124], "廢胎費規則：KF 單價內含廢胎費，需於業績計算時扣減；KD 單價外加廢胎費，不需扣減。")
append_highlight(doc.paragraphs[124], "【待確認／會議討論】需確認 KF/KD 與公司別 1/F 對應關係，並與 DOC-01/DOC-06 廢胎費規則一致。")
set_paragraph_text(doc.paragraphs[126], "公司別／胎別判斷：")
set_paragraph_text(doc.paragraphs[127], "前台登入後需依經銷商對應之公司別＋胎別，決定可見功能、人機介面載入內容與語系顯示。")
set_paragraph_text(doc.paragraphs[128], "此判斷邏輯應由後端依 dealer/company_code/product_series 等資料回傳前台，不由使用者人工選擇。")
set_paragraph_text(doc.paragraphs[130], "架構需保留多公司別、多胎別與多語系擴充彈性；是否納入本期實作範圍另行確認。")

erp_apis = doc.tables[14]
set_cell_text(erp_apis.rows[5].cells[3], "多年月／起迄年月獎金數字；分類呈現包含固定、販促、其它獎金。")

# Security / token uniqueness.
auth = doc.tables[17]
set_cell_text(auth.rows[1].cells[1], "登入後取得 Bearer token，存於 ir.config_parameter，以 kenda.token.<公司別>.<客代> 管控；本案以「公司別＋客代」作為唯一識別值。")
set_cell_text(auth.rows[2].cells[0], "X-Company-Code／X-Customer-Code Header")
set_cell_text(auth.rows[2].cells[1], "每次 API 請求須帶公司別與客代，後端雙重驗證 token、company_code、customer_code 一致。")

data_sec = doc.tables[20]
set_cell_text(data_sec.rows[3].cells[1], "kenda.erp.sync.log 記錄每次 ERP API 呼叫（request_json / response_json / state / retry）；授權人員可查詢錯誤通知與完整同步記錄。")

perf = doc.tables[21]
set_cell_text(perf.rows[5].cells[1], "由 ir.cron 定期拉取並暫存於 kenda.performance/bonus/reconciliation，前台查詢不即時呼叫 ERP；獎金分類與顯示方式需與業務規則確認。")

# Future evolution / to-do.
todo = doc.tables[26]
set_cell_text(todo.rows[3].cells[0], "P2")
set_cell_text(todo.rows[3].cells[1], "V4 多公司別全球語系架構預留")
set_cell_text(todo.rows[3].cells[2], "現行規劃需預留多公司別、全球語系與介面切換彈性。")
append_highlight_to_cell(todo.rows[3].cells[2], "【待確認／會議討論】需確認此項僅為架構預留，或需本期實作；若需實作，應評估範圍與工時。")
row = todo.add_row()
for idx, value in enumerate(["P2", "廣告資產管理範圍確認", "確認「廣告資產申請及查詢／各式廣告資產合約下載」是否為店招功能延伸或新增模組。"]):
    set_cell_text(row.cells[idx], value)
append_highlight_to_cell(row.cells[2], "【待確認／會議討論】若屬新增模組，需確認是否在原合約範圍內。")

# v1.1 summary cleanup.
set_paragraph_text(doc.paragraphs[261], "結帳自動拆單：購物車「確認送出」時，後端依「報價單編號＋包裝部門」規則，自動拆分為多張建大訂單（系統自動取號），並於前台顯示「拆單結果」。")
set_paragraph_text(doc.paragraphs[262], "本案無三層報價功能需求，並已取消「廢胎費（含／不含）」作為拆單層級；廢胎費規則改由公司別／接單公司對應之業績計算邏輯處理。")
set_paragraph_text(doc.paragraphs[267], "廣告資產／店招流程為獨立流程：廣告資產申請、合約下載與經銷商下單為不同流程，不可混用。")
set_paragraph_text(doc.paragraphs[268], "若廣告資產管理超出既有店招申請與合約下載範圍，需另行確認是否列入本案。")

patch_header_footer_versions(doc)
doc.save(OUT)

# Patch version text inside footer drawing/textbox XML when present.
tmp = OUT.with_suffix(".tmp.docx")
with zipfile.ZipFile(OUT, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename.startswith("word/footer") and item.filename.endswith(".xml"):
            data = data.replace(b"v1.0", b"v3.0")
        zout.writestr(item, data)
shutil.move(str(tmp), str(OUT))
print(str(OUT))
