from copy import deepcopy
from pathlib import Path
import shutil
import zipfile

from docx import Document
from docx.enum.text import WD_COLOR_INDEX


SRC = Path(r"C:\Users\jenny.lu\Documents\建大輪胎專案\期中交付文件\DOC-01專案範圍與需求規格書2.0.docx")
OUT = Path(r"C:\Users\jenny.lu\Documents\建大輪胎專案\期中交付文件\DOC-01專案範圍與需求規格書3.0.docx")


def clone_run_style(src, dst):
    dst.bold = src.bold
    dst.italic = src.italic
    dst.underline = src.underline
    if src.font is not None:
        dst.font.name = src.font.name
        dst.font.size = src.font.size
        dst.font.color.rgb = src.font.color.rgb


def set_paragraph_text(paragraph, text):
    base = paragraph.runs[0] if paragraph.runs else None
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    if base is not None:
        clone_run_style(base, run)


def set_cell_text(cell, text):
    paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, text)
    for extra in cell.paragraphs[1:]:
        set_paragraph_text(extra, "")


def append_highlight(paragraph, text):
    if paragraph.text and not paragraph.text.endswith((" ", "\n")):
        paragraph.add_run(" ")
    run = paragraph.add_run(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True
    return run


def append_highlight_to_cell(cell, text):
    paragraph = cell.paragraphs[-1]
    if paragraph.text:
        paragraph.add_run("\n")
    append_highlight(paragraph, text)


def clear_row(row):
    for cell in row.cells:
        set_cell_text(cell, "")


def delete_table(table):
    tbl = table._element
    tbl.getparent().remove(tbl)


def rebuild_appendix_table(doc):
    # Rebuild the appendix table because the source table is split into many
    # continuation rows, which makes local edits render as duplicate fragments.
    app1 = doc.tables[35]
    app2 = doc.tables[36]
    delete_table(app2)
    tbl = app1._tbl
    for tr in list(tbl.tr_lst)[1:]:
        tbl.remove(tr)

    rows = [
        ["A-01", "§5 關係人", "ox 許啟裕單位應為「營業五部」", "已更正為「營業五部」（§5、§11.3）。", "✓ 已修訂"],
        ["A-02", "§2.1 組織背景", "產品缺機車胎（MC）、轎車胎（PCR）、電動車胎（EB/EM/EV）；品牌為「建大輪胎」。", "產品系列改列 BC/MC/PCR/IC/EB-EM-EV，並標註品牌建大輪胎；各公司別實際系列以產品主檔為準。", "✓ 已修訂"],
        ["A-03", "§2.2 現況", "現況為「3 個 Website 下單或依賴電話」。", "已補述須透過 3 個不同網站分別下單或依賴電話／email。", "✓ 已修訂"],
        ["A-04", "§3／§8.2", "廢胎費屬生產廠出貨代收代付；公司別=1 外加不扣、公司別=F 內含要扣。", "已於名詞定義與業績計算更正為依公司別計算，並補建大／建豐兩範例。", "✓ 已修訂"],
        ["A-05", "§3 名詞", "建大天津廠（Tianjin）、內胎/無內胎、TPI（簾布密度）等產品屬性。", "已補充建大天津廠、內胎/無內胎及 TPI（簾布密度）等產品規格欄位；實際廠別與產品屬性以建大提供之產品主檔為準。", "✓ 已修訂"],
        ["A-06", "§4.3／§8", "沒印象談及三層報價功能。", "已確認本案無三層報價功能需求；結帳自動拆單邏輯為「報價單編號＋包裝部門」，三層報價如需導入應列為後續變更或新需求。", "✓ 已澄清"],
        ["A-07", "§4.3／v1.1", "取消廢胎費做為拆單邏輯；目前只以報價單編號＋包裝部門作為拆單依據。", "文件拆單規則已改為二層（報價單編號＋包裝部門），取消廢胎費層。", "✓ 已修訂"],
        ["A-08", "§5 關係人", "主擔當：Bohaw 林柏豪。", "已新增「專案主擔當 Bohaw 林柏豪」。", "✓ 已修訂"],
        ["A-09", "§5 關係人", "本案應無聘請外部顧問。", "已移除「外部顧問（建達顧問）」並加註本案無外部顧問。", "✓ 已修訂"],
        ["A-10", "§6 權限", "下單／產品目錄／訂單查詢需依角色隱藏部份資訊。", "已新增 §6.4「欄位層級權限」，明定單價／廢胎費／包裝依角色與帳號權限顯示與否。", "✓ 已修訂"],
        ["A-11", "§6.2 矩陣", "階別 4（經銷商業務）理論上可申請店招。", "功能矩陣「店招申請」已開放階別 4。", "✓ 已修訂"],
        ["A-12", "§6／§11", "海外廠沒有廢胎費，要預留顯示與否選項。", "已於 §3、§6.4 註明海外廠無廢胎費、需預留顯示選項。", "✓ 已修訂"],
        ["A-13", "§6", "登入後人機介面要同步載入該客代所屬胎別（產品系列）介面。", "已於 §6.4 與 DOC-02 商品總覽說明「登入即依客代所屬產品系列載入，不需人工挑選」。", "✓ 已修訂"],
        ["A-14", "FR-017", "依客代所屬廠別上載不同語系版本的合約清單。", "FR-017 已補「依廠別顯示對應語系合約清單」。", "✓ 已修訂"],
        ["A-15", "§7／§8", "多處「缺公司別」「缺公司別+客代」。", "已補「公司別＋客代＝唯一值」，並於名詞定義、API 說明與業績規則補充公司別欄位。", "✓ 已修訂"],
        ["A-16", "FR-024／§8.1", "轉單失敗需提示錯誤訊息。", "FR-024 已補「轉單失敗須提示明確錯誤訊息，使操作人員清楚知悉訂單轉單失敗原因」。", "✓ 已修訂"],
        ["A-17", "NF-P04", "若要支援同時上線 500/1000 位不降速，需要哪些資源？", "本案目標 100 併發；500/1000 併發需擴充 Odoo workers、CPU/RAM、PostgreSQL 連線池調校、前置負載平衡與靜態 CDN 快取。正式規格將另提資源評估。", "▣ 待另提評估"],
        ["A-18", "§11.3", "能否在 Odoo 上查看完整記錄？", "Odoo 後台保有訂單、拆單群組、ERP 同步日誌等完整記錄。\n【待確認／會議討論】需確認可查看完整記錄之對象、權限層級與操作畫面。", "■ 待確認"],
        ["A-19", "§7／§8", "特殊報價單（伙伴雲－免費樣品單）需由業務或助理在平台下單轉單。", "【待確認／會議討論】已於報價單查詢／下單相關需求加螢光標示，會議確認是否納入本版範圍。", "■ 待確認"],
        ["A-20", "§10／§11", "需搭配 KENDA BPM 表單流程進簽測試，才能 UAT。", "已補充 BPM 表單流程及串接測試為 UAT/Go-Live 前置相依，若流程未完成可能影響 UAT 時程。", "✓ 已修訂"],
        ["A-21", "§11", "要求加入後續持續優化直至符合需求之文字。", "【待確認／會議討論】不納入開放式承諾；建議於會議回報此類文字可能造成驗收標準不明確與修改範圍無上限。", "■ 待確認"],
        ["A-22", "§6／§8", "購物車及訂單查詢若選擇不顯示金額時，亦不得顯示總金額。", "已於欄位層級權限補充螢光提示，待確認畫面與 API 欄位遮蔽方式。", "■ 待確認"],
        ["A-23", "§8／§11", "各胎別顯示模式需可設定，例如 BC/MC 以金額、PCR 以銷售量計算達成率。", "已於權限／業績規則／驗收標準加螢光標示，待確認功能控制點與設定方式。", "■ 待確認"],
    ]
    for values in rows:
        row = app1.add_row()
        for idx, value in enumerate(values):
            set_cell_text(row.cells[idx], value)
            if value.startswith("【待確認") or "\n【待確認" in value:
                for paragraph in row.cells[idx].paragraphs:
                    for run in paragraph.runs:
                        if "【待確認" in run.text or "會議討論" in run.text:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                            run.bold = True


def replace_in_paragraph(paragraph, old, new):
    if old not in paragraph.text:
        return False
    set_paragraph_text(paragraph, paragraph.text.replace(old, new))
    return True


def add_table_row_after(table, after_idx, values):
    src_row = table.rows[after_idx]
    tr = deepcopy(src_row._tr)
    table._tbl.insert(after_idx + 1, tr)
    new_row = table.rows[after_idx + 1]
    for i, value in enumerate(values):
        if i < len(new_row.cells):
            set_cell_text(new_row.cells[i], value)
    return new_row


doc = Document(SRC)

# Cover / metadata
replace_in_paragraph(doc.paragraphs[3], "KENDA B2B v1.2", "KENDA B2B v3.0")
replace_in_paragraph(doc.paragraphs[26], "文件版本\tv1.2", "文件版本\tv3.0")
set_paragraph_text(doc.paragraphs[27], "發行日期\t2026-06-22")
replace_in_paragraph(doc.paragraphs[32], "v1.1 → v1.2，2026-06-09", "v1.2 → v3.0，2026-06-22")
replace_in_paragraph(doc.paragraphs[264], "v1.1 → v1.2，2026-06-09", "v1.2 → v3.0，2026-06-22")
for section in doc.sections:
    for part in (section.header, section.footer, section.first_page_header, section.first_page_footer):
        for paragraph in part.paragraphs:
            replace_in_paragraph(paragraph, "v1.2", "v3.0")
        for table in part.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, "v1.2", "v3.0")

# Revision history
rev = doc.tables[2]
row = rev.add_row()
for idx, value in enumerate(["v3.0", "2026-\n06-22", "艾創點", "依 0618 期中待修正文件與修訂清單更新：補充產品屬性、BPM 串接相依、伙伴雲/查詢/達成率待確認提示，並修正拆單、權限、欄位與術語。"]):
    set_cell_text(row.cells[idx], value)

# Organization and terms
replace_in_paragraph(doc.paragraphs[82], "公司別代碼：\t", "公司別代碼：1")
replace_in_paragraph(doc.paragraphs[83], "公司別代碼：\t", "公司別代碼：F")

terms_1 = doc.tables[6]
set_cell_text(terms_1.rows[4].cells[2], "經銷商識別資料之一；本案以「公司別＋客代」作為唯一識別值，例：1+R02019、F+M01011。")
terms_2 = doc.tables[7]
set_cell_text(terms_2.rows[5].cells[1], "Kenda China / Kenda Tianjin / Kenda Thailand")
set_cell_text(terms_2.rows[5].cells[2], "建大中國廠／建大天津廠／泰國廠（生產廠別依建大產品主檔為準）")
set_cell_text(terms_2.rows[7].cells[2], "內胎型 / 無內胎型")
add_table_row_after(terms_2, 8, ["TPI", "Threads Per Inch", "簾布密度，作為產品規格欄位之一；實際欄位值以建大產品主檔／規格資料為準。"])

# Scope and out-of-scope wording
scope_out = doc.tables[9]
set_cell_text(scope_out.rows[5].cells[0], "三層報價功能")
set_cell_text(scope_out.rows[5].cells[1], "本案無三層報價功能需求。結帳自動拆單邏輯已確認為「報價單編號＋包裝部門」，此項與經銷商階層定價結構不同，三層報價如需導入應列為後續變更或新需求。")

# Roles / permissions
roles = doc.tables[11]
set_cell_text(roles.rows[3].cells[2], "大型經銷商主帳號，可查看自身及二級代理商／下線群組資料。")
set_cell_text(roles.rows[4].cells[2], "一級代理的財務／業務／採購帳號，可查看自身授權範圍內資料；不得查看二級代理商資訊。")

field_perm = doc.tables[13]
set_cell_text(field_perm.rows[1].cells[1], "一級經銷商（主帳號）向建大取得之單價")
set_cell_text(field_perm.rows[3].cells[1], "包裝部門／包裝說明")

replace_in_paragraph(doc.paragraphs[151], "一組／主帳號", "一級／主帳號")
replace_in_paragraph(doc.paragraphs[153], "包裝資訊", "包裝部門／包裝資訊")
append_highlight(doc.paragraphs[153], "【待確認／會議討論】購物車及訂單查詢若依權限設定為不顯示金額，畫面亦不得顯示總金額。")
append_highlight(doc.paragraphs[154], "【待確認／會議討論】各胎別顯示模式／達成率計算模式可能需可設定，例如 BC/MC 以金額、PCR 以銷售量計算達成率；此功能應再確認設定控制點與系統範圍。")

# Functional requirement details
front = doc.tables[15]
append_highlight_to_cell(front.rows[3].cells[3], "【待確認／會議討論】特殊報價單（伙伴雲－免費樣品單）是否需由業務或助理於平台下單並轉單，需於會議確認是否納入本版範圍。")
set_cell_text(front.rows[9].cells[3], "依年/月查詢；主帳號可帶下線 dealer 參數查群組明細+總計。二級代理商資料僅代理商主帳號可查看，一級代理之財務／業務／採購帳號不得查看二級代理商資訊。")

back = doc.tables[17]
set_cell_text(back.rows[5].cells[2], "確認後拋轉建大 ERP（state: confirmed → pushed → acked）；轉單失敗時須於後台提示明確錯誤訊息，使操作人員清楚知悉訂單轉單失敗原因。")

erp = doc.tables[19]
set_cell_text(erp.rows[2].cells[4], "依公司別＋客代取得最新報價單頭／單身（含單價、廢胎費、包裝部門、內胎/無內胎、TPI 等規格欄位；實際欄位以 ERP／產品主檔為準）")
set_cell_text(erp.rows[4].cells[4], "依公司別＋客代＋年月清單查詢業績總額（邏輯在 ERP，只回最終值）")

perf = doc.tables[20]
set_cell_text(perf.rows[7].cells[1], "經銷商以「公司別＋客代」作為唯一識別值。")
append_highlight_to_cell(perf.rows[9].cells[1], "【待確認／會議討論】各胎別達成率計算基準（金額／銷售量）是否需由後台設定，需確認功能控制點。")

# Business constraints and acceptance
constraints = doc.tables[29]
set_cell_text(constraints.rows[3].cells[1], "退貨流程、三層報價功能等複雜功能不含於本案範圍，列入後續版本或另行需求變更。")
add_table_row_after(constraints, 5, ["BC-06", "本案 UAT 需搭配 KENDA BPM 表單流程及相關串接測試確認；若 BPM 表單流程或簽核節點未完成，可能影響 UAT 時程與驗收排程。"])

uat_table = doc.tables[31]
set_cell_text(uat_table.rows[7].cells[1], "業績金額依 ERP 回傳邏輯與公司別廢胎費規則計算，與 ERP 一致。")
append_highlight_to_cell(uat_table.rows[7].cells[1], "【待確認／會議討論】若各胎別達成率計算基準不同（金額／銷售量），需確認設定方式與驗收案例。")

replace_in_paragraph(doc.paragraphs[227], "建⼤營業五部", "建大營業五部")
set_paragraph_text(doc.paragraphs[228], "建大資訊部（Cliff）驗證 ERP API、KENDA BPM 表單流程及相關串接測試正確性")
set_paragraph_text(doc.paragraphs[229], "建大業務／資訊人員確認業績、獎金與流程測試結果")
append_highlight(doc.paragraphs[230], "【待確認／會議討論】UAT 須搭配 KENDA BPM 表單流程進行串接測試，待 BPM 流程測試可執行後再確認完整 UAT 結果。")

golive = doc.tables[34]
add_table_row_after(golive, 1, ["KENDA BPM 表單流程及串接測試完成，並可支援 UAT 驗證", "建大資訊部 / 艾創點"])
set_cell_text(golive.rows[6].cells[0], "操作手冊完成並完成「經銷商網路下單平台_教育訓練說明會」")

# v1.1 summary split paragraphs
set_paragraph_text(doc.paragraphs[252], "結帳自動拆單（依客戶意見修正為二層）：購物車「確認送出」時，後端依序以 ①報價單編號 ②包裝部門 規則，自動拆分為多張建大訂單（系統自動取號），並於前台顯示「拆單結果」。")
set_paragraph_text(doc.paragraphs[253], "依客戶意見，已取消「廢胎費（含／不含）」作為拆單層級（廢胎費含／不含由公司別決定，不另作拆單依據）。")

# Appendix A updates
app1 = doc.tables[35]
set_cell_text(app1.rows[13].cells[0], "A-05")
set_cell_text(app1.rows[13].cells[1], "§3 名詞")
set_cell_text(app1.rows[13].cells[2], "建大天津廠（Tianjin）、內胎/無內胎、TPI（簾布密度）等產品屬性。")
set_cell_text(app1.rows[13].cells[3], "已補充建大天津廠、內胎/無內胎及 TPI（簾布密度）等產品規格欄位；實際廠別與產品屬性以建大提供之產品主檔為準。")
set_cell_text(app1.rows[13].cells[4], "✅ 已修訂")
clear_row(app1.rows[14])
clear_row(app1.rows[15])
set_cell_text(app1.rows[16].cells[3], "已確認本案無三層報價功能需求；結帳自動拆單邏輯為「報價單編號＋包裝部門」，三層報價如需導入應列為後續變更或新需求。")
set_cell_text(app1.rows[20].cells[0], "A-07")
set_cell_text(app1.rows[20].cells[1], "§4.3／v1.1")
set_cell_text(app1.rows[20].cells[2], "取消廢胎費做為拆單邏輯；目前只以報價單編號＋包裝部門作為拆單依據。")
set_cell_text(app1.rows[20].cells[3], "文件拆單規則已改為二層（報價單編號＋包裝部門），取消廢胎費層。")
set_cell_text(app1.rows[20].cells[4], "✅ 已修訂")
clear_row(app1.rows[21])
clear_row(app1.rows[22])
clear_row(app1.rows[23])

app2 = doc.tables[36]
set_cell_text(app2.rows[9].cells[3], "已補「公司別＋客代＝唯一值」，並於名詞定義、API 說明與業績規則補充公司別欄位。")
set_cell_text(app2.rows[11].cells[3], "FR-024 已補「轉單失敗須提示明確錯誤訊息，使操作人員清楚知悉訂單轉單失敗原因」。")
set_cell_text(app2.rows[18].cells[3], "Odoo 後台保有訂單、拆單群組、ERP 同步日誌等完整記錄。")
append_highlight_to_cell(app2.rows[18].cells[3], "【待確認／會議討論】需確認可查看完整記錄之對象、權限層級與操作畫面。")
set_cell_text(app2.rows[18].cells[4], "🟨 待確認")

# Add new appendix rows for 0618 notes that should remain visible.
new_rows = [
    ["A-19", "§7/§8", "特殊報價單（伙伴雲－免費樣品單）需由業務或助理在平台下單轉單。", "【待確認／會議討論】已於報價單查詢／下單相關需求加螢光標示，會議確認是否納入本版範圍。", "🟨 待確認"],
    ["A-20", "§10/§11", "需搭配 KENDA BPM 表單流程進簽測試，才能 UAT。", "已補充 BPM 表單流程及串接測試為 UAT/Go-Live 前置相依，若流程未完成可能影響 UAT 時程。", "✅ 已修訂"],
    ["A-21", "§11", "要求加入後續持續優化直至符合需求之文字。", "【待確認／會議討論】不納入開放式承諾；建議於會議回報此類文字可能造成驗收標準不明確與修改範圍無上限。", "🟨 待確認"],
    ["A-22", "§6/§8", "購物車及訂單查詢若選擇不顯示金額時，亦不得顯示總金額。", "已於欄位層級權限補充螢光提示，待確認畫面與 API 欄位遮蔽方式。", "🟨 待確認"],
    ["A-23", "§8/§11", "各胎別顯示模式需可設定，例如 BC/MC 以金額、PCR 以銷售量計算達成率。", "已於權限/業績規則/驗收標準加螢光標示，待確認功能控制點與設定方式。", "🟨 待確認"],
]
for values in new_rows:
    row = app2.add_row()
    for idx, value in enumerate(values):
        set_cell_text(row.cells[idx], value)
        if value.startswith("【待確認"):
            row.cells[idx].paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW

rebuild_appendix_table(doc)

doc.save(OUT)

# Footer version sits inside a footer drawing/textbox in the source document,
# so patch footer XML after saving.
tmp = OUT.with_suffix(".tmp.docx")
with zipfile.ZipFile(OUT, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename.startswith("word/footer") and item.filename.endswith(".xml"):
            data = data.replace(b"v1.2", b"v3.0")
        zout.writestr(item, data)
shutil.move(str(tmp), str(OUT))
print(str(OUT))
