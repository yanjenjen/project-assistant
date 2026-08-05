from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / '00_系統說明與戰略' / '個人成長體系手冊_2026_V1.0.docx'
NAVY='17365D'; BLUE='2F75B5'; PALE='D9EAF7'; LIGHT='F2F4F7'; GOLD='FFF2CC'; GRAY='666666'; WHITE='FFFFFF'

doc=Document()
sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11)
sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
sec.header_distance=sec.footer_distance=Inches(0.492)

styles=doc.styles
normal=styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11); normal.font.color.rgb=RGBColor.from_string('1F2937')
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for name,size,color,before,after in [('Title',28,NAVY,0,10),('Heading 1',16,BLUE,18,10),('Heading 2',13,BLUE,14,7),('Heading 3',12,NAVY,10,5)]:
    st=styles[name]; st.font.name='Calibri'; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color)
    st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
for lname in ['List Bullet','List Number']:
    st=styles[lname]; st.font.name='Calibri'; st.font.size=Pt(11); st.paragraph_format.left_indent=Inches(.375); st.paragraph_format.first_line_indent=Inches(-.188); st.paragraph_format.space_after=Pt(4); st.paragraph_format.line_spacing=1.25

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)
def margins(cell,top=80,start=120,bottom=80,end=120):
    tc=cell._tc.get_or_add_tcPr(); tcMar=tc.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tc.append(tcMar)
    for tag,val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn('w:'+tag))
        if node is None: node=OxmlElement('w:'+tag); tcMar.append(node)
        node.set(qn('w:w'),str(val)); node.set(qn('w:type'),'dxa')
def set_table_widths(table,widths):
    table.autofit=False; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    tblPr=table._tbl.tblPr; tblW=tblPr.find(qn('w:tblW'))
    if tblW is None: tblW=OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'),str(sum(widths))); tblW.set(qn('w:type'),'dxa')
    tblInd=tblPr.find(qn('w:tblInd'))
    if tblInd is None: tblInd=OxmlElement('w:tblInd'); tblPr.append(tblInd)
    tblInd.set(qn('w:w'),'120'); tblInd.set(qn('w:type'),'dxa')
    grid=table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        gc=OxmlElement('w:gridCol'); gc.set(qn('w:w'),str(w)); grid.append(gc)
    for row in table.rows:
        for cell,w in zip(row.cells,widths):
            cell.width=Inches(w/1440); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cell)
            tcW=cell._tc.get_or_add_tcPr().find(qn('w:tcW')); tcW.set(qn('w:w'),str(w)); tcW.set(qn('w:type'),'dxa')
def add_table(headers, rows, widths):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; set_table_widths(t,widths)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c,BLUE)
        for r in c.paragraphs[0].runs: r.font.bold=True; r.font.color.rgb=RGBColor.from_string(WHITE); r.font.size=Pt(9.5)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=1.1
                for r in p.runs: r.font.size=Pt(9.5)
    set_table_widths(t,widths); doc.add_paragraph().paragraph_format.space_after=Pt(1); return t
def callout(label,text,fill=GOLD):
    t=doc.add_table(rows=1,cols=1); t.style='Table Grid'; set_table_widths(t,[9360]); shade(t.cell(0,0),fill)
    p=t.cell(0,0).paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(label+'：'); r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY); p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
def bullet(text): doc.add_paragraph(text,style='List Bullet')
def page_break(): doc.add_page_break()

# Header/footer
hp=sec.header.paragraphs[0]; hp.text='PGOS｜個人成長作業系統'; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs: r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(GRAY)
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
fp.add_run('2026 V1.0  |  ')
fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)

# Cover
p=doc.add_paragraph(); p.style='Title'; p.alignment=WD_ALIGN_PARAGRAPH.LEFT; p.add_run('個人成長作業系統')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT; r=p.add_run('Personal Growth Operating System（PGOS）'); r.font.size=Pt(16); r.font.color.rgb=RGBColor.from_string(BLUE)
doc.add_paragraph('ERP 智慧流程顧問長期成長制度｜2026 V1.0')
doc.add_paragraph('\n')
add_table(['文件欄位','內容'],[['版本','2026 V1.0'],['狀態','試運轉'],['建置日期','2026-08-04'],['系統所有人','Jenny Lu'],['檢討週期','每週、每月、每季、每年'],['正式資料位置','個人成長體系資料夾（Git 追蹤）']],[2700,6660])
callout('核心原則','對話不是系統；可持續更新、可攜、可稽核的檔案才是系統。能力升級必須有真實證據。',PALE)
page_break()

doc.add_heading('1. 制度目的與適用範圍',level=1)
doc.add_paragraph('本制度用於長期管理職涯方向、能力發展、思維與軟實力、實戰專案、成果證據、知識資產及市場價值。其目的不是增加行政負擔，而是使日常 ERP 顧問工作能持續轉化為更高價值、可移轉及可被證明的能力。')
doc.add_heading('1.1 暫定三年定位',level=2)
callout('定位假設','成為能運用 AI、資料與流程方法，協助企業診斷 ERP 營運問題、設計改善方案並推動落地的智慧流程顧問。')
doc.add_paragraph('此定位不是永久承諾。第一次季度能力考核時，應依本人意願、工作機會、實際成果與 AI 市場變化重新校準。')
doc.add_heading('1.2 範圍與排除事項',level=2)
add_table(['納入','不納入'],[['職涯策略、七大能力、實戰專案','無使用情境的課程收藏'],['案例、作品、方法、模板與AI工作流','未去識別化客戶原始資料'],['外部回饋、週月季檢討、身心狀態','不必要的私人醫療與敏感資料'],['市場成果及職涯選擇權','只為美觀而維護的複雜頁面']],[4680,4680])

doc.add_heading('2. PGOS 系統架構',level=1)
add_table(['層級','管理問題','正式產出'],[['方向層','我要往哪裡走？','三年定位、年度主題、選擇原則'],['能力層','需要形成什麼能力？','能力字典、現況、目標、差距'],['專案層','如何取得真實練習？','效率型、能力型、資產型專案'],['執行層','每週如何推進？','12週計畫、每週成果與時間'],['證據層','如何證明真的變強？','案例、作品、回饋及結果'],['治理層','如何避免偏航？','週、月、季、年度檢討']],[1500,3200,4660])
doc.add_heading('2.1 工具分工',level=2)
add_table(['工具','用途','使用原則'],[['Excel管理中心','盤點、計畫、指標、圖表','結構化摘要及連結，不存長篇內容'],['Word／Markdown','制度、策略、方法、檢討','Git可追蹤內容優先保留Markdown'],['資料夾','案例、證據、模板、作品','固定命名、去識別化'],['行事曆／自動提醒','讓檢討確實發生','週、月、季固定節奏'],['AI教練','提問、整理、挑戰及檢查','不能代替本人自評與真實驗證']],[1800,3100,4460])

page_break()
doc.add_heading('3. 七大能力資本',level=1)
add_table(['領域','發展重點','代表能力'],[['思維與決策','辨識、推理、取捨及復盤','系統思考、問題定義、因果、機率、批判思考'],['ERP與流程','企業流程、資料、制度與系統','跨模組、主資料、內控、異常、產業情境'],['資料與AI','資料分析、AI、整合與自動化','Excel、SQL、資料品質、API、AI評測'],['顧問能力','從需求到落地','訪談、診斷、方案、測試、變更管理'],['軟實力','溝通、關係、協作與影響','傾聽、表達、談判、衝突、信任'],['執行與身心','支撐長期執行','專注、壓力調節、韌性、持續學習'],['職涯與市場','把能力轉成選擇權','作品、方法、關係網絡、可移轉能力']],[1800,3100,4460])
doc.add_heading('3.1 思維與軟實力的培養方式',level=2)
doc.add_paragraph('思維與軟實力不能靠閱讀或單次自評形成，必須放進具體情境，透過行為、結果、回饋與復盤觀察。')
add_table(['循環','實施方式'],[['情境','選擇真實決策、訪談、衝突、簡報或協作事件'],['行為','事前設定一項可觀察行為，如先重述對方觀點再回應'],['結果','記錄是否改善理解、承諾、決策或關係'],['回饋','向主管、同事、客戶或AI取得具體觀察'],['復盤','辨識假設、偏誤、逃避或過度承諾'],['再練習','在下一個相似情境刻意調整一項行為']],[1800,7560])
callout('禁止事項','不得只用「溝通能力3分」判定軟實力；必須附情境、行為、結果及信心程度。')

doc.add_heading('4. 評量與證據規則',level=1)
add_table(['等級','名稱','行為判定'],[['L0','尚未理解','無法辨識或說明'],['L1','能說明','可說明概念，尚無法穩定使用'],['L2','能協作完成','在模板、AI或他人提醒下完成'],['L3','能獨立完成','一般真實情境可獨立完成'],['L4','能處理例外','複雜、高壓或例外情境仍有效'],['L5','能建立方法','可教導他人並持續改善方法']],[1100,1700,6560])
add_table(['證據','定義','是否可支持升級'],[['E0 無證據','純自我感覺','否'],['E1 知識','筆記、測驗、概念說明','僅支持L1'],['E2 模擬','練習或模擬案例','原則上不超過L2'],['E3 實戰','真實專案交付物','可支持L3'],['E4 結果','時間、品質、採用或外部成果','可支持L3–L5，仍需多次情境']],[1500,5500,2360])
doc.add_heading('4.1 能力考核原則',level=2)
bullet('沒有 E3 真實實戰，原則上不得升為 L3。')
bullet('L4 必須有多次複雜或例外情境，不能只靠一次成功。')
bullet('能力分數旁必須標示低、中、高證據信心。')
bullet('外部回饋是校準工具，不是把成長主導權交給他人。')
bullet('分數是導航，不是人格價值或績效懲罰。')

page_break()
doc.add_heading('5. 成長專案與學習機制',level=1)
doc.add_paragraph('實戰專案是成長的主引擎。課程、閱讀及工具練習只在支援當前專案時列入核心計畫。')
add_table(['類型','目的','ERP顧問例子'],[['效率型','釋放低價值時間','訪談紀錄轉需求、SOP初稿、文件一致性檢查'],['能力型','發展一項尚未成熟能力','獨立主持訪談、SQL差異查核、設計例外流程'],['資產型','把一次工作轉成複利資產','題庫、原因樹、檢核表、方法論、AI工作流'],['複合型','同時產生效率、能力及資產','建立AI輔助需求分析流程並在專案驗證']],[1500,2300,5560])
doc.add_heading('5.1 專案選擇',level=2)
doc.add_paragraph('候選專案以工作頻率、工作價值、未來定位、可形成證據、可重用及八週可完成六項各1–5分評估。優先選擇高頻、高價值、可重用且規模可控的項目。')
doc.add_heading('5.2 學習停止規則',level=2)
bullet('四週內沒有實際使用場景。')
bullet('無法說明它支援哪項定位或專案。')
bullet('只有輸入，沒有成果、練習或回饋。')
bullet('只是因為熱門而學，或維護成本高於實際效益。')
bullet('AI已能完成，而本人沒有必要理解底層或承擔判斷。')

doc.add_heading('6. 12週執行制度',level=1)
add_table(['規劃欄位','限制'],[['核心成果','一項，必須可驗收'],['核心能力','一項'],['支援能力','一至二項'],['實戰專案','一個主要專案'],['可重用資產','一項'],['每週成果','可看見、可完成、能形成證據']],[2200,7160])
doc.add_heading('6.1 每週最小循環',level=2)
add_table(['步驟','行動'],[['選定成果','以交付物描述，不寫成「學習某主題」'],['安排時間','將實戰、必要學習及沉澱放入行事曆'],['完成最小成果','形成查核表、流程圖、查詢、訪談、案例或工具'],['取得回饋','用實際結果、外部觀察或AI反例挑戰'],['沉澱資產','至少留下案例、方法、模板或AI工作流之一']],[1800,7560])
doc.add_heading('6.2 建議每週三小時最低配置',level=2)
add_table(['時間','用途'],[['90分鐘','實戰專案'],['45分鐘','依專案需求學習'],['30分鐘','整理案例、方法或模板'],['15分鐘','週檢討及下週成果']],[1800,7560])

page_break()
doc.add_heading('7. 知識與資產管理',level=1)
add_table(['資產庫','保存內容','最低必要欄位'],[['案例庫','真實問題與結果','背景、假設、查證、方案、結果、復盤'],['問題模式庫','跨案例共通模式','症狀、可能原因、查證、控制點'],['方法與模板','已驗證做法','適用範圍、步驟、例外、版本'],['AI工作流','人機協作流程','輸入、工具、人工點、測試、風險、成效'],['回饋與評量','外部觀察','情境、具體行為、影響、建議'],['能力證據','支持能力等級的資料','能力、E等級、連結、去識別化']],[1700,3000,4660])
doc.add_heading('7.1 Git 與資訊安全',level=2)
bullet('客戶名稱、聯絡人、帳號、個資、密碼、真實交易資料及商業機密不得進入公開或一般Git版本。')
bullet('案例採代號並去識別化；大型或敏感附件保存在核准位置，PGOS只存索引與摘要。')
bullet('提交前檢查檔名、文件屬性、註解、隱藏欄與圖片是否洩漏資訊。')

doc.add_heading('8. 治理與長期營運',level=1)
add_table(['頻率','時間上限','必做事項'],[['每週','20分鐘','成果、能力、證據、障礙、下週唯一重點'],['每月','60分鐘','時間配置、案例沉澱、實戰比例、繼續停止開始'],['每季','3小時','證據審查、能力校準、策略調整、下一個12週計畫'],['每年','半天','三年定位、年度成果、市場價值、次年主題與封存']],[1300,1500,6560])
doc.add_heading('8.1 角色責任',level=2)
add_table(['角色','責任'],[['本人','提供真實資料、執行、判斷、確認評分及資料安全'],['AI教練','提問、整理、反方挑戰、檢查一致性、協助更新正式檔案'],['外部回饋者','針對指定成果提供具體觀察，不代替本人決策']],[1800,7560])

doc.add_heading('9. 儀表板與量化原則',level=1)
doc.add_paragraph('量化用來發現方向與偏差，不追求虛假的精確。思維與軟實力以季度行為證據為主；執行、成果與資產可用週月數據觀察。')
add_table(['指標類型','代表指標','用途'],[['投入','實際時數、實戰／學習比例','檢查資源是否投入'],['產出','案例、證據、資產、專案數','檢查是否留下成果'],['能力','獨立完成、例外處理、外部回饋','判定能力是否升級'],['結果','工時、品質、採用、客戶／職涯成果','確認是否轉化為價值'],['狀態','能量、壓力、阻礙','避免短期衝刺破壞長期營運']],[1500,4200,3660])
callout('圖表解讀','雷達或長條圖只看能力分布；趨勢圖看投入與成果。沒有證據的分數不得作為升級結論。',PALE)

page_break()
doc.add_heading('10. 第一個90天建置與試運轉',level=1)
add_table(['期間','工作','驗收成果'],[['第1–2週','完成戰略、任務及初次能力盤點','三年定位假設、現況分數與證據信心'],['第3–4週','選擇核心能力及首個成長專案','12週計畫、行事曆時間、外部回饋者'],['第5–8週','執行專案並每週留下最小成果','至少2筆E3候選證據及1次外部回饋'],['第9–10週','將經驗形成資產','1份案例、1個模板／方法／AI工作流'],['第11–12週','正式季度能力考核','能力校準、專案成果、下一季計畫']],[1500,4300,3560])
doc.add_heading('10.1 第一輪啟動清單',level=2)
bullet('填寫Excel「成長戰略」中的待盤點欄位。')
bullet('完成七大能力現況分數；每個非零分數至少寫一則證據摘要。')
bullet('從真實工作列出至少三個候選成長專案並評分。')
bullet('只選一個本季執行專案，設定12週核心成果。')
bullet('指定至少一位可以提供具體回饋的人。')
bullet('確認週、月、季提醒時間是否適合。')

doc.add_heading('11. 仍待本人確認事項',level=1)
add_table(['待確認','首次完成時點'],[['三年定位是否符合本人真正意願','第一次月檢討前'],['優先深耕流程／模組','第一個12週計畫前'],['每週可穩定投入時間','啟動當週'],['外部回饋者','第4週前'],['提醒時間是否需要調整','第一次提醒後'],['Git為公開或私人倉庫及敏感資料政策','提交PGOS前']],[6400,2960])
callout('完成定義','V1.0的完成不是所有欄位都已個人化，而是制度、工具、節奏與第一輪入口均已可運作。第一次12週循環負責校準。')

# metadata and save
doc.core_properties.title='個人成長作業系統（PGOS）2026 V1.0'
doc.core_properties.subject='ERP智慧流程顧問長期成長制度'
doc.core_properties.author='Jenny Lu'
OUT.parent.mkdir(parents=True,exist_ok=True)
doc.save(OUT)
print(OUT)
