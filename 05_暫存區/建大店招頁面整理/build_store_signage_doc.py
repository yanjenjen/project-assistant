from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "建大工業_經銷商平台專案" / "200_專案資料" / "240_功能需求管理" / "features" / "F005_店招申請"
DOCX_PATH = OUT_DIR / "06_功能開發與調整整理_20260708.docx"
MD_PATH = OUT_DIR / "06_功能開發與調整整理_20260708.md"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(table)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], "F2F4F7")
        for p in hdr[i].paragraphs:
            p.runs[0].bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    set_table_widths(table, widths)
    doc.add_paragraph()
    return table


def add_kv_table(doc, rows):
    return add_table(doc, ["項目", "內容"], rows, [1.65, 4.85])


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def add_note(doc, title, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    for line in lines:
        p = cell.add_paragraph(style=None)
        p.add_run(line)
    set_table_widths(table, [6.5])
    doc.add_paragraph()


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def build_docx():
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("建大系統－店招頁面功能開發與調整整理")
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("功能編號：F005 店招申請｜整理日期：2026-07-08｜來源：店招畫面.pptx").italic = True

    doc.add_heading("一、文件目的與適用範圍", level=1)
    doc.add_paragraph(
        "本文件依據建大提供之「店招畫面.pptx」整理店招頁面的功能開發與調整需求，"
        "供需求確認、系統設計、開發排程與驗收測試使用。內容涵蓋經銷商入口網站前台、"
        "後台維護作業、Excel 上下載資料流、BPM／建大資訊處理邊界，以及 2026-06-02 測試後需修正或補入的項目。"
    )
    add_kv_table(
        doc,
        [
            ("適用系統", "建大工業股份有限公司－四輪網路下單系統（https://b2b.kenda.com.tw/cdr/EDTB/pcr/）"),
            ("功能範圍", "店招製作申請、施工項目維護、國家／省縣維護、合約書下載、進度查詢、合約到期查詢。"),
            ("不在艾創點範圍", "廣告編號產生、BPM 簽核流程、申請書 PDF 產生、工作委託單作業。"),
            ("主要使用角色", "經銷商／客代、各廠商企人員、總公司商企人員、建大資訊、建大業務。"),
        ],
    )

    doc.add_heading("二、整體角色與權限原則", level=1)
    add_table(
        doc,
        ["角色", "主要作業", "權限重點"],
        [
            ("經銷商／客代前台使用者", "提出店招申請、下載合約書、查詢申請進度與合約到期資料。", "僅可查看登入帳號所屬國家、廠別、胎種別及自身申請範圍內資料。申請人、國家由登入帳號自動帶出，不可手填或修改。"),
            ("各廠商企人員", "維護該廠店招資料、下載前台申請資料、更新 Excel 後上載。", "一般權限限於所屬廠別資料，可查詢、下載、上載；更新與刪除權限依建大授權控管。"),
            ("總公司商企人員", "跨廠資料維護、欄位開放設定、施工項目與後台主檔維護。", "最高權限，可新增、修改、刪除、查詢、下載、上載，並控制哪些欄位開放前台顯示。"),
            ("建大資訊", "廣告資產編號取號、BPM 簽核、申請書 PDF 產出。", "廣告編號依 BPM 結案與系統規則產生，避免跨廠／海外廠流水號重複。"),
            ("建大業務", "工作委託單及內部業務申請作業。", "工作委託單不在本入口網站／艾創點作業範圍。"),
        ],
        [1.25, 2.55, 2.7],
    )
    add_note(
        doc,
        "資料隔離原則",
        [
            "前台查詢與下載不得跨經銷商、跨國家或跨胎種別顯示資料。",
            "後台各廠人員僅能維護所屬廠資料，避免上載 Excel 時覆蓋其他廠資料。",
            "總公司商企人員保留跨廠最高維護權限，供集團同步資料與例外修正使用。",
        ],
    )

    doc.add_heading("三、功能模組整理", level=1)
    modules = [
        ("F005-1", "店招製作申請", "前台申請表單，包含新增、維修、作廢移招；續約、盤點、作廢資產轉移先預留後台開關，未來視流程調整是否開放。"),
        ("F005-1A", "施工項目類別維護", "後台維護申請類別、廠別、胎種別、是否開放前台及更新日期，總公司商企需具最大維護權限。"),
        ("F005-1B", "施工項目維護", "後台維護店級別、施工項目細項、廠別、胎種別、自填選項與是否開放前台。"),
        ("F005-1C", "國家／省縣維護", "後台維護國家、省／縣市及適用廠別，前台依登入帳號國家提供下拉選單。"),
        ("F005-2", "店招合約書下載", "前台依登入帳號之國家、廠別、胎種別、店級別顯示可下載合約書；後台維護合約書檔案。"),
        ("F005-4", "店招進度查詢", "前台以廣告編號、車行名稱、地址或經銷商模糊查詢自身申請進度；後台透過 Excel 下載、補齊、上載維護查詢資料。"),
        ("F005-3", "店招合約到期查詢", "前台查詢自身範圍內合約到期資料，支援合約日期區間與 PDF 下載；後台與進度查詢共用 Excel 資料源。"),
    ]
    add_table(doc, ["編號", "模組", "整理後功能定位"], modules, [0.8, 1.7, 4.0])

    doc.add_heading("四、前台功能需求", level=1)
    doc.add_heading("4.1 店招製作申請", level=2)
    add_bullet(doc, "申請人與國家依登入帳號自動帶出，不可自行填入或修改。")
    add_bullet(doc, "廣告物所在縣市以國家篩選下拉選單提供，縣市清單由後台維護。")
    add_bullet(doc, "需補入廣告物編號欄位；若已申請過，可輸入廣告編號查詢並自動帶出歷史資料。")
    add_bullet(doc, "車行名稱、聯絡人、電話、地址等欄位可由申請者填寫；若透過廣告編號帶出歷史資料，需可預填。")
    add_bullet(doc, "施工項目需包含大項、店級別、細項與自填欄位，且分類項次需符合建大提供設定。")
    add_bullet(doc, "需補入合約書起始日、合約書到期日欄位。")
    add_bullet(doc, "畫面外觀需重新整理，避免欄位散亂，並保留未來開放續約、盤點、作廢資產轉移的彈性。")

    doc.add_heading("4.2 合約書下載", level=2)
    add_bullet(doc, "前台不顯示廠別／胎種別欄位，由登入帳號設定自動判斷可下載範圍。")
    add_bullet(doc, "只顯示該帳號所屬國家、廠別、胎種別與店級別可用的合約書範本。")
    add_bullet(doc, "合約清單需可點擊下載，顯示欄位至少包含店級別、更新日期、檔名。")

    doc.add_heading("4.3 進度查詢", level=2)
    add_bullet(doc, "前台僅能查詢該客代帳號曾申請範圍內資料，不可查其他客代帳號資料。")
    add_bullet(doc, "查詢條件需支援廣告編號、車行名稱、地址或經銷商名稱，並支援模糊搜尋。")
    add_bullet(doc, "前台僅開放查詢 Excel 中申請狀態為新增、維修、作廢移招之資料；盤點、續約、作廢合約轉移不開放前台查詢。")
    add_bullet(doc, "列表欄位建議至少包含申請進度、車行、申請項目、更新日期；需預留欄位開關與欄位名稱調整彈性。")
    add_bullet(doc, "顯示筆數需分頁，例如每頁 10 筆，避免畫面過長。")

    doc.add_heading("4.4 合約到期查詢", level=2)
    add_bullet(doc, "前台依登入經銷商顯示該經銷商現有申請資料之合約到期清單。")
    add_bullet(doc, "查詢條件需支援車行名稱、廣告編號、經銷商、地址，以及合約起訖日區間。")
    add_bullet(doc, "查詢結果需可下載為 PDF。")
    add_bullet(doc, "列表顯示建議每頁最多 10 筆，超出可下一頁。")

    doc.add_heading("五、後台維護與資料流需求", level=1)
    add_table(
        doc,
        ["後台功能", "維護欄位／資料", "作業重點"],
        [
            ("施工項目類別維護", "申請類別、廠別、胎種別、是否開放前台、更新日期。", "同廠別資料需排序在一起；總公司商企最大權限，可跨廠維護。"),
            ("施工項目維護", "項目、店級別、廠別、胎種別、是否開放前台、更新日期。", "需支援自填選項，例如胎種別、車體彩繪、一般店等。"),
            ("國家／省縣維護", "國家、省／縣市、廠別。", "前台依登入帳號國家顯示下拉選單；可評估由建大客戶維護主檔撈取或於後台維護。"),
            ("合約書維護", "廠別、胎種別、店級別、上載資料、更新日期、上載人員。", "需支援上載、刪除、下載、修改；更新日期由程式依上載日自動填入。"),
            ("進度查詢資料維護", "前台申請資料下載、人工補齊欄位、Excel 上載更新。", "下載檔名需含查詢區間、廠別、胎種別、日期；上載前提醒資料將覆蓋。"),
            ("前台欄位開放設定", "Excel 欄位是否開放前台顯示、欄位名稱。", "保留未來欄位增減與名稱調整彈性，由授權後台人員控制。"),
        ],
        [1.55, 2.25, 2.7],
    )

    doc.add_heading("六、廣告編號與 BPM 邊界", level=1)
    add_bullet(doc, "廣告資產編號由建大資訊處理，於經銷商入口網站申請後進入 BPM，主辦送出並經主管簽核同意後才產出。")
    add_bullet(doc, "廣告編號原則：第一碼為項目分類，第二、三碼為西元年末二碼，後四碼為流水號。例：2026 年形象店 F260001；2026 年直招 A260001。")
    add_bullet(doc, "分類開頭：F 為形象店，A 為直／橫招，C 為車體彩繪。")
    add_bullet(doc, "因海外廠也可能共用程式，需避免同時申請造成流水號重複。")
    add_bullet(doc, "申請書 PDF 由建大資訊依既有 SCH010 欄位資料產生，檔名最後 7 碼為廣告物資產編號。")
    add_note(
        doc,
        "待確認",
        [
            "流水號是否每年重置、跨廠是否共用流水號、同號衝突如何鎖定。",
            "入口網站是否需與 BPM 回寫廣告編號、回寫時點與欄位格式。",
            "續約、盤點、作廢資產轉移未來若由 BPM 提出，如何同步至入口網站 Excel 查詢資料源。",
        ],
    )

    doc.add_heading("七、2026-06-02 測試後調整清單", level=1)
    add_table(
        doc,
        ["項目", "測試後需調整內容", "影響模組"],
        [
            ("申請書前台", "廣告物編號、施工項目大項、施工項目分類／自填欄位、合約書起訖日未完整設計，需補入。", "F005-1"),
            ("縣市下拉", "廣告物縣市需以下拉式選單供申請人選擇，並依登入國家篩選。", "F005-1／F005-1C"),
            ("申請人／國家", "需依登入帳號自動帶出，不可自行輸入或修改。", "F005-1"),
            ("後台維護", "施工項目類別、施工項目、國家省縣維護在測試畫面未看到，需補入或確認資料來源。", "F005-1A／1B／1C"),
            ("合約書下載", "前台測試有誤，需依登入帳號廠別／胎種別過濾，不顯示其他廠別或胎種別資料。", "F005-2"),
            ("合約後台", "後台合約書上載、刪除、下載、更新資料功能測試未看到，需補入。", "F005-2"),
            ("進度查詢前台", "需補入廣告編號、車行名稱、地址、經銷商等模糊查詢條件，並限定該客代曾申請資料。", "F005-4"),
            ("進度查詢後台", "需補入區間下載、Excel 上載、下載最新版 Excel、欄位開放設定與覆蓋提醒。", "F005-4"),
            ("到期列表前台", "需補入合約起訖日期區間查詢、欄位與 PDF 下載功能。", "F005-3"),
        ],
        [1.25, 3.7, 1.55],
    )

    doc.add_heading("八、驗收重點", level=1)
    checks = [
        "前台登入後，申請人、國家、廠別、胎種別判斷正確，且不得手動修改應由系統帶出的欄位。",
        "經銷商前台不可查看其他經銷商、其他國家、其他胎種別或其他廠別資料。",
        "店招申請表單欄位完整，包含廣告物編號、縣市下拉、施工項目大項與細項、自填欄位、合約書起訖日。",
        "廣告物編號查詢可帶出歷史資料，且新增申請未有廣告編號時允許空白送出。",
        "施工項目類別、施工項目、國家／省縣、合約書檔案可由後台維護，並依權限限制資料範圍。",
        "合約書下載清單依登入帳號範圍過濾，非該帳號適用資料不顯示。",
        "進度查詢支援指定條件與模糊搜尋，僅顯示新增、維修、作廢移招等對前台開放狀態。",
        "進度與到期資料 Excel 下載／上載流程可用，檔名規則、覆蓋提醒、欄位開放設定完整。",
        "合約到期查詢支援合約起訖日期區間與 PDF 下載。",
        "列表畫面具分頁或筆數控制，避免畫面過長。",
    ]
    for item in checks:
        add_bullet(doc, item)

    doc.add_heading("九、待確認事項", level=1)
    add_table(
        doc,
        ["編號", "待確認事項", "建議處理方向"],
        [
            ("Q1", "廣告編號完整產生規則、流水號鎖定與跨廠／海外廠排重機制。", "由建大資訊提供 BPM 取號規格，確認入口網站是否需接收回寫。"),
            ("Q2", "國家／省縣資料來源是否由建大客戶維護主檔撈取，或仍需建立入口網站後台維護功能。", "若主檔穩定可串接，後台僅保留查詢與例外維護；若不穩定則需完整 CRUD 與 Excel 維護。"),
            ("Q3", "續約、盤點、作廢資產轉移未來從 BPM 申請後，如何進入進度／到期查詢資料源。", "建議確認 BPM 是否可輸出或 API 回寫，避免人工逐筆補 Excel。"),
            ("Q4", "Excel 上載筆數限制與效能，來源提到單一廠可能已有 4 萬多筆資料。", "需進行匯入效能與檔案大小限制評估，必要時改批次匯入或背景處理。"),
            ("Q5", "前台欄位開放設定的顆粒度：欄位顯示、欄位名稱、排序、是否可下載是否都需後台控制。", "建議先定義最小可行欄位清單，預留欄位名稱與顯示開關。"),
        ],
        [0.55, 3.1, 2.85],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("F005 店招申請｜功能開發與調整整理｜2026-07-08")

    doc.save(DOCX_PATH)


def build_markdown():
    md = """# 【F005】店招頁面功能開發與調整整理

- 功能編號：F005 店招申請
- 整理日期：2026-07-08
- 來源：`店招畫面.pptx`
- 適用系統：建大工業股份有限公司－四輪網路下單系統
- 目的：整理店招頁面功能開發、測試後修正與待確認事項，供需求確認、開發與驗收使用。

## 一、功能範圍

| 編號 | 模組 | 功能定位 |
|---|---|---|
| F005-1 | 店招製作申請 | 前台申請表單，包含新增、維修、作廢移招；續約、盤點、作廢資產轉移先預留後台開關。 |
| F005-1A | 施工項目類別維護 | 後台維護申請類別、廠別、胎種別、是否開放前台及更新日期。 |
| F005-1B | 施工項目維護 | 後台維護店級別、施工項目細項、廠別、胎種別、自填選項與是否開放前台。 |
| F005-1C | 國家／省縣維護 | 後台維護國家、省／縣市及適用廠別，前台依登入帳號國家提供下拉選單。 |
| F005-2 | 店招合約書下載 | 前台依登入帳號之國家、廠別、胎種別、店級別顯示可下載合約書；後台維護合約書檔案。 |
| F005-4 | 店招進度查詢 | 前台以廣告編號、車行名稱、地址或經銷商模糊查詢自身申請進度；後台透過 Excel 維護查詢資料。 |
| F005-3 | 店招合約到期查詢 | 前台查詢自身範圍內合約到期資料，支援日期區間與 PDF 下載；後台與進度查詢共用 Excel 資料源。 |

## 二、角色與權限

| 角色 | 主要作業 | 權限重點 |
|---|---|---|
| 經銷商／客代前台使用者 | 提出申請、下載合約書、查詢進度與到期資料 | 僅可查看登入帳號所屬國家、廠別、胎種別及自身申請範圍內資料。 |
| 各廠商企人員 | 維護該廠資料、下載前台申請資料、更新 Excel 後上載 | 一般權限限所屬廠別資料。 |
| 總公司商企人員 | 跨廠資料維護、欄位開放設定、主檔維護 | 最高權限，可新增、修改、刪除、查詢、下載、上載與控制前台欄位顯示。 |
| 建大資訊 | 廣告資產編號取號、BPM 簽核、申請書 PDF 產出 | 廣告編號由 BPM 簽核後產出。 |
| 建大業務 | 工作委託單及內部業務申請作業 | 工作委託單不在艾創點作業範圍。 |

## 三、前台需求重點

### 3.1 店招製作申請

- 申請人與國家依登入帳號自動帶出，不可自行填入或修改。
- 廣告物所在縣市以國家篩選下拉選單提供，縣市清單由後台維護。
- 需補入廣告物編號欄位；若已申請過，可輸入廣告編號查詢並自動帶出歷史資料。
- 施工項目需包含大項、店級別、細項與自填欄位，且分類項次需符合建大提供設定。
- 需補入合約書起始日、合約書到期日欄位。
- 保留未來開放續約、盤點、作廢資產轉移的彈性。

### 3.2 合約書下載

- 前台不顯示廠別／胎種別欄位，由登入帳號設定自動判斷可下載範圍。
- 只顯示該帳號所屬國家、廠別、胎種別與店級別可用的合約書範本。
- 合約清單需可點擊下載，顯示欄位至少包含店級別、更新日期、檔名。

### 3.3 進度查詢

- 僅能查詢該客代帳號曾申請範圍內資料，不可查其他客代帳號資料。
- 查詢條件需支援廣告編號、車行名稱、地址或經銷商名稱，並支援模糊搜尋。
- 前台僅開放查詢申請狀態為新增、維修、作廢移招之資料。
- 列表欄位建議至少包含申請進度、車行、申請項目、更新日期；需預留欄位開關與欄位名稱調整彈性。
- 顯示筆數需分頁，例如每頁 10 筆。

### 3.4 合約到期查詢

- 依登入經銷商顯示該經銷商現有申請資料之合約到期清單。
- 查詢條件需支援車行名稱、廣告編號、經銷商、地址，以及合約起訖日區間。
- 查詢結果需可下載為 PDF。
- 列表建議每頁最多 10 筆，超出可下一頁。

## 四、後台維護與資料流

| 後台功能 | 維護欄位／資料 | 作業重點 |
|---|---|---|
| 施工項目類別維護 | 申請類別、廠別、胎種別、是否開放前台、更新日期 | 同廠別資料需排序；總公司商企最大權限。 |
| 施工項目維護 | 項目、店級別、廠別、胎種別、是否開放前台、更新日期 | 需支援自填選項。 |
| 國家／省縣維護 | 國家、省／縣市、廠別 | 前台依登入國家顯示下拉；需確認資料來源。 |
| 合約書維護 | 廠別、胎種別、店級別、上載資料、更新日期、上載人員 | 支援上載、刪除、下載、修改。 |
| 進度查詢資料維護 | 前台申請資料下載、人工補齊欄位、Excel 上載更新 | 下載檔名需含查詢區間、廠別、胎種別、日期；上載前提醒覆蓋。 |
| 前台欄位開放設定 | Excel 欄位是否開放前台顯示、欄位名稱 | 保留欄位增減與名稱調整彈性。 |

## 五、2026-06-02 測試後調整清單

| 項目 | 需調整內容 | 影響模組 |
|---|---|---|
| 申請書前台 | 廣告物編號、施工項目大項、施工項目分類／自填欄位、合約書起訖日未完整設計，需補入。 | F005-1 |
| 縣市下拉 | 廣告物縣市需以下拉式選單供申請人選擇，並依登入國家篩選。 | F005-1／F005-1C |
| 申請人／國家 | 需依登入帳號自動帶出，不可自行輸入或修改。 | F005-1 |
| 後台維護 | 施工項目類別、施工項目、國家省縣維護在測試畫面未看到，需補入或確認資料來源。 | F005-1A／1B／1C |
| 合約書下載 | 前台需依登入帳號廠別／胎種別過濾，不顯示其他廠別或胎種別資料。 | F005-2 |
| 合約後台 | 合約書上載、刪除、下載、更新資料功能需補入。 | F005-2 |
| 進度查詢前台 | 需補入廣告編號、車行名稱、地址、經銷商等模糊查詢條件，並限定該客代曾申請資料。 | F005-4 |
| 進度查詢後台 | 需補入區間下載、Excel 上載、下載最新版 Excel、欄位開放設定與覆蓋提醒。 | F005-4 |
| 到期列表前台 | 需補入合約起訖日期區間查詢、欄位與 PDF 下載功能。 | F005-3 |

## 六、待確認事項

| 編號 | 待確認事項 | 建議處理方向 |
|---|---|---|
| Q1 | 廣告編號完整產生規則、流水號鎖定與跨廠／海外廠排重機制。 | 由建大資訊提供 BPM 取號規格，確認入口網站是否需接收回寫。 |
| Q2 | 國家／省縣資料來源是否由建大客戶維護主檔撈取，或仍需建立入口網站後台維護功能。 | 若主檔穩定可串接，後台僅保留查詢與例外維護；若不穩定則需完整 CRUD 與 Excel 維護。 |
| Q3 | 續約、盤點、作廢資產轉移未來從 BPM 申請後，如何進入進度／到期查詢資料源。 | 建議確認 BPM 是否可輸出或 API 回寫，避免人工逐筆補 Excel。 |
| Q4 | Excel 上載筆數限制與效能，來源提到單一廠可能已有 4 萬多筆資料。 | 需進行匯入效能與檔案大小限制評估，必要時改批次匯入或背景處理。 |
| Q5 | 前台欄位開放設定的顆粒度。 | 建議先定義最小可行欄位清單，預留欄位名稱與顯示開關。 |
"""
    MD_PATH.write_text(md, encoding="utf-8")


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_markdown()
    print(DOCX_PATH)
    print(MD_PATH)
