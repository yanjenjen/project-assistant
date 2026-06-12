from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SRC = r"C:\Users\jenny.lu\Documents\建大輪胎專案\艾創點-輔導備忘錄260519.docx"
OUT = r"C:\Users\jenny.lu\Documents\艾創點數位-ERP顧問\05_暫存區\艾創點-輔導備忘錄260519_追繳文件版.docx"


def clear_cell(cell):
    for paragraph in list(cell.paragraphs):
        p = paragraph._element
        p.getparent().remove(p)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_paragraph(cell, text="", bold=False, size=10, space_after=2):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_bullet(cell, text, level=0):
    p = add_paragraph(cell, "", size=10, space_after=1)
    p.paragraph_format.left_indent = Pt(18 + level * 12)
    run = p.add_run("• " + text)
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(10)


doc = Document(SRC)
table = doc.tables[0]

table.rows[0].cells[1].text = "建大經銷商下單平台-第二次會議（追繳文件整理版）"

content_cell = table.rows[5].cells[0]
clear_cell(content_cell)

add_paragraph(content_cell, "1. 問題：【AI 開發所需資料與測試環境】", bold=True)
add_paragraph(
    content_cell,
    "開發團隊需取得實際測試資料，包含客戶資料、報價單、訂單及相關欄位，作為 AI 判讀、自動填寫與系統邏輯建立依據。目前資料仍未完整，將影響開發與驗證進度。",
)
add_paragraph(content_cell, "1.1 解決方案：", bold=True)
add_bullet(content_cell, "客戶端預計於 2026/05/22 前整理測試資料，優先提供報價單與訂單相關欄位及實際資料。")
add_bullet(content_cell, "資料結構建議以 Excel 方式提供，便於欄位比對、資料清理與後續匯入。")
add_bullet(content_cell, "顧問端收到資料後，先區分「可建檔」、「待補正」、「待釐清」三類，避免分批補件造成遺漏。")

add_paragraph(content_cell, "1.2 待確認事項：", bold=True)
add_bullet(content_cell, "確認 API 權限申請流程、申請人、核准單位與預計完成日，並於內部群組回報進度。")
add_bullet(content_cell, "確認目前已比對完成且可先行輸入系統之 4 個檔案／4 個商品清單，需補上商品名稱、檔名、版本與來源。")

add_paragraph(content_cell, "2. 問題：【系統介面與規格產出】", bold=True)
add_paragraph(
    content_cell,
    "本次會議討論商品列表、購物車、報價單、合約與設施等畫面之初步版本。後續系統規格、程式開發、測試錄影與操作手冊需同步產出，確保功能與文件一致。",
)
add_paragraph(content_cell, "2.1 解決方案：", bold=True)
add_bullet(content_cell, "開發流程採「規格 -> AI 寫 Code -> 自動測試／錄影 -> 自動產出操作手冊」方式推進。")
add_bullet(content_cell, "使用者若對畫面有調整意見，建議以截圖標註方式回饋，便於快速修改前端畫面。")
add_bullet(content_cell, "系統需依登入身分進行權限分流，例如客戶組長、產品組長可見單據與功能應有所區分。")

add_paragraph(content_cell, "2.2 待確認事項：", bold=True)
add_bullet(content_cell, "針對 VC 流程中較複雜之系列產品規格呈現方式，需再由客戶端確認規格邏輯與畫面呈現規則。")

add_paragraph(content_cell, "3. 問題：【資料串接與編碼規則】", bold=True)
add_paragraph(
    content_cell,
    "為確保後續資料整合準確，並降低表單、欄位與來源資料串接錯誤，需統一各項代碼、欄位與流水號規則。",
)
add_paragraph(content_cell, "3.1 解決方案：", bold=True)
add_bullet(content_cell, "客戶端需提供「胎種別」、「產品別」等專有名詞之中英文對照代碼表。")
add_bullet(content_cell, "客戶端需提供現行流水號、來源代碼或申請規則之生成原則，供開發團隊寫入系統邏輯。")

add_paragraph(content_cell, "4. 代辦事項：【追繳文件】", bold=True)
for item in [
    "追繳報價單與訂單相關欄位及實際測試資料；預計提供日為 2026/05/22，仍需確認實際交付狀態。",
    "追繳客戶資料、商品主檔、商品規格、分類資料及其他系統建檔必要欄位。",
    "追繳「胎種別」、「產品別」等代碼表，需含中文名稱、英文名稱、代碼、適用範圍與備註。",
    "追繳現行流水號、來源代碼、申請規則或版本規則，作為資料串接與系統判斷依據。",
    "確認已可先行輸入系統之 4 個檔案／4 個商品，並建立清單追蹤建檔狀態。",
    "確認 API 權限申請進度，包含申請窗口、審核窗口、權限範圍與預計開通日期。",
    "確認會議錄影、側錄或字幕檔保存位置，作為後續補正會議紀錄與追蹤決議依據。",
]:
    add_bullet(content_cell, item)

add_paragraph(content_cell, "5. 後續會議與追蹤方式", bold=True)
add_bullet(content_cell, "建議建立「追繳文件清單」，欄位包含文件名稱、對應商品／流程、負責窗口、預計提供日、實際收到日、文件版本、處理狀態與備註。")
add_bullet(content_cell, "顧問端每次收到補件後，需回覆是否可建檔、是否需補正，以及是否影響開發或測試排程。")
add_bullet(content_cell, "下次會議應優先檢視追繳文件清單，確認未結項目、逾期項目與需客戶決策事項。")

set_cell_shading(content_cell, "FFFFFF")

doc.save(OUT)
print(OUT)
